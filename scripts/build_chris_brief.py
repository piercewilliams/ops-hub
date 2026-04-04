"""
Build: Chris Palo Huddle Brief (2026-04-03)
Run: python3 scripts/build_chris_brief.py
Output: ~/Desktop/Chris Palo Huddle Brief - 2026-04-03.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.expanduser("~/Desktop/Chris Palo Huddle Brief - 2026-04-03.docx")

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

# ── Style helpers ─────────────────────────────────────────────────────────────
DARK  = RGBColor(0x1a, 0x1a, 0x2e)
BLUE  = RGBColor(0x16, 0x47, 0x8a)
TEAL  = RGBColor(0x0f, 0x72, 0x6e)
GRAY  = RGBColor(0x55, 0x55, 0x66)
WHITE = RGBColor(0xff, 0xff, 0xff)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    run.font.name = "Calibri"
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '164789')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEAL
    run.font.name = "Calibri"
    return p

def body(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.name = "Calibri"
        r1.font.color.rgb = DARK
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.name = "Calibri"
    r2.font.color.rgb = DARK
    return p

def bullet(text, bold_prefix=None, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.25 + indent * 0.2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + "  ")
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.name = "Calibri"
        r1.font.color.rgb = DARK
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.name = "Calibri"
    r2.font.color.rgb = DARK
    return p

def callout(text):
    """Shaded callout box via paragraph shading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EAF2FB')
    pPr.append(shd)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK
    return p

def spacer(pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ═══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after  = Pt(2)
r = title_p.add_run("Chris Palo Huddle — Key Takeaways")
r.bold = True
r.font.size = Pt(16)
r.font.name = "Calibri"
r.font.color.rgb = DARK

sub_p = doc.add_paragraph()
sub_p.paragraph_format.space_before = Pt(0)
sub_p.paragraph_format.space_after  = Pt(14)
rs = sub_p.add_run("April 3, 2026  ·  2:54–3:33 PM  ·  Pierce Williams & Chris Palo")
rs.font.size = Pt(10)
rs.font.name = "Calibri"
rs.font.color.rgb = GRAY

callout(
    "This is a synthesis of every strategic concept, framework, goal, priority, and implicit expectation "
    "that surfaced in the April 3 huddle — including things that may not have registered fully in the moment. "
    "Read it as a briefing, not just notes."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 1. SYNDICATION ECOSYSTEM TAXONOMY
# ═══════════════════════════════════════════════════════════════════════════════
h1("1. Syndication Ecosystem Taxonomy — Two Completely Separate Worlds")

body(
    "Chris draws a hard line between two types of syndication environments. These are not just different "
    "platforms — they are different strategic games with different rules. Mixing data from both produces "
    "conclusions that are meaningless at best, actively misleading at worst."
)

h2("App-Based Captured Environments")
bullet("Apple News", bold_prefix="Platforms:")
bullet("SmartNews")
bullet("Newsbreak")
body(
    "Users are inside an app. They never leave the app. They will never become McClatchy subscribers. "
    "The audience is 'captured' in the sense that they're already inside a closed ecosystem — McClatchy's content "
    "is surfaced to them by the platform's algorithm. The goal is purely to maximize impressions and page views "
    "within that environment. There is no downstream subscriber or LTV play. The competition is: can McClatchy "
    "content beat other publishers for algorithm attention inside the app?"
)

h2("Web-Based Competitive Environments")
bullet("Yahoo / Yahoo News", bold_prefix="Platforms:")
bullet("O&O (McClatchy owned-and-operated sites)")
body(
    "Users land on a real webpage. They can be retargeted. They can subscribe. Discovery dynamics apply — "
    "the article has to earn a click from a search result or feed, then hold the reader long enough to matter. "
    "Standard CTR, PV, and engagement metrics apply here in the way you'd expect."
)

callout(
    "Why this matters for your work: Any headline formula analysis that averages Apple News results with Yahoo "
    "results is comparing two different sports. A formula that dominates in Apple News (captured, algorithm-surfaced) "
    "may perform terribly in Yahoo (competitive, click-from-feed). The T1 analysis engine needs to respect this "
    "boundary — run comparisons within ecosystem type only."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 2. LTV = 0 FOR SYNDICATION
# ═══════════════════════════════════════════════════════════════════════════════
h1("2. LTV = 0 for Syndication Content")

body(
    "Every article that goes out on Apple News, SmartNews, Yahoo, etc. also lives on O&O. "
    "The syndication version has a Lifetime Value of zero — no subscriber conversion, no retargeting, "
    "no downstream revenue relationship. The only value it produces is incremental page views."
)

body(
    "This is not a criticism — it's a framework. It means syndication strategy should be framed entirely "
    "around: which platform, which content type, and which headline formula produces the most incremental PV "
    "delta on top of what the article would get organically on O&O. Every optimization decision in the "
    "syndication layer flows from this."
)

callout(
    "PRD implication: When writing about why syndication content strategy matters, the framing is not "
    "'grow the audience' or 'build readership' — it's 'maximize PV increments across platforms that never convert.' "
    "That framing also clarifies why headline formula research (P3) is so valuable: a better formula = more "
    "incremental PVs across every article McClatchy syndicates."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. CLUSTER BATTING AVERAGE
# ═══════════════════════════════════════════════════════════════════════════════
h1("3. Cluster Batting Average — How Chris Actually Measures Success")

body(
    "This is Chris's primary performance metric for the CSA program. It measures how often a story "
    "outperforms average PVs across multiple platforms simultaneously — the 'double and triple dip.'"
)

h2("The Metric")
bullet("Definition: a story that performs above average PVs on 2 or more platforms")
bullet("Q1 goal: 1-in-4 articles hits this mark")
bullet("Current rate: ~1-in-3 (above goal)")
bullet("Before CSA: 1-in-5 hit rate, with far fewer articles in the pipeline")
bullet("Now: 1-in-3 hit rate with ~5x more articles going through the system")

h2("Why 'Double and Triple Dipping' Is the Core Value Prop")
body(
    "The reason this metric matters is that the same article hitting on Apple News AND SmartNews AND Yahoo "
    "produces three separate PV streams that each contribute independently. Each platform hit is additive. "
    "A story that 'triple dips' is generating 3x the PV value of a story that only performs on one platform. "
    "This is what Chris means when he talks about the multiplier effect of syndication."
)

callout(
    "This metric is what Chris brings to Chris Tarot and division leadership. When he says the CSA program "
    "is working, this is the evidence he uses. Understanding this number and its trend line is critical for "
    "anyone reporting upward on this initiative."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 4. THREE TEST TYPES
# ═══════════════════════════════════════════════════════════════════════════════
h1("4. Three Test Types — Chris's Testing Taxonomy")

body(
    "Chris has a specific vocabulary for content experiments. These are not interchangeable. "
    "Understanding which type you're running matters for how you design it, measure it, and report it."
)

h2("Format Testing")
body(
    "A/B testing headline formulas, structures, or presentation approaches against each other. "
    "Example: does a 'Here's why' formula outperform a question headline on Apple News? "
    "This is what the T1 analysis engine primarily produces. Fixed content, variable format."
)

h2("Content Testing")
body(
    "Testing topics and story angles — does a particular type of content (travel, local, nostalgia) "
    "outperform across platforms? Sarah Price's testing tracker is full of content tests. "
    "Variable content, possibly fixed format. This maps to Sarah's current Q1 testing work "
    "(Nostalgia, Swarming, Travel/Experiences, Everyday Life tracks)."
)

h2("Swarm Testing")
body(
    "The most sophisticated type. A performing story triggers a follow-up cluster — multiple variants "
    "that 'swarm' the topic from different angles while it's hot. The manual version of this exists "
    "right now in Sarah Price's testing tracker (5 story clusters, 21 articles, P/C tagging). "
    "Chris's vision is to automate this: a story performs → CSA detects it → CSA recommends next cluster "
    "topic → writers execute → cluster performs → feedback loop updates governor."
)

callout(
    "Swarm testing is the north star for what the CSA system should eventually do automatically. "
    "The Cluster ID infrastructure currently in dev (PGS-40) is a prerequisite for automating it at scale. "
    "The manual swarming test in Sarah's tracker is evidence of demand — use it as the re-pitch for P7 "
    "when the infrastructure lands."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 5. CONTENT AS INVENTORY
# ═══════════════════════════════════════════════════════════════════════════════
h1("5. Content as Inventory — The Mental Model Chris Uses")

body(
    "Every slot on every McClatchy website is inventory. Like a retailer's shelf space, it can be "
    "filled with owned product (McClatchy original content) or partner/vendor product (United Robots, "
    "Reuters, wire services). The goal is to maximize the performance of every slot."
)

body(
    "This mental model has three implications:"
)

bullet(
    "Owned content is always preferable to partner content — higher margin, full editorial control, "
    "reusable in the syndication layer. Partner content fills gaps but at a cost."
)
bullet(
    "The United Robots initiative (P12) is really an inventory capture play — McClatchy generating "
    "its own automated alert-based content instead of paying United Robots for theirs."
)
bullet(
    "The partner content optimization (P15) is an inventory audit — figure out which slots are "
    "being filled by partner content and whether owned content could fill them more profitably."
)

callout(
    "Chris mentioned that a single focused audit of partner content inventory could produce a ~10% traffic "
    "improvement in one week. He said 'not yet' — meaning he knows what the opportunity is, he's choosing "
    "the timing. Watch for the signal that he's ready to activate this."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 6. PARTNER CONTENT / P15
# ═══════════════════════════════════════════════════════════════════════════════
h1("6. Partner Content / Inventory Optimization (P15) — Not Yet, But It's Coming")

body(
    "Chris raised this explicitly. He has a clear picture of the opportunity and a quantified estimate "
    "(~10% traffic lift from a one-time optimization pass). He is choosing when to activate it, not whether."
)

bullet("Kathy owns the partner content relationship — she must be in the room when this starts.")
bullet("Reuters is being evaluated as a potential replacement for the current partner content provider.")
bullet("No action until Chris green-lights. But be ready — this could be activated quickly once he says go.")

# ═══════════════════════════════════════════════════════════════════════════════
# 7. POLITICAL DATA — TWO SEPARATE WORLDS
# ═══════════════════════════════════════════════════════════════════════════════
h1("7. Political Data — Two Separate Worlds, Don't Reconcile Them")

body(
    "There are two entirely separate data worlds at McClatchy, and Chris is clear that they should "
    "not be conflated or forced to reconcile:"
)

h2("World 1: Macro Political Performance Dashboards")
body(
    "Justin's and Dedra's dashboards — big-picture political content performance numbers, division-level "
    "tracking, executive-facing metrics. This is where the 'whatabouter' dynamic lives (political content "
    "that triggers high engagement via outrage/reaction but doesn't convert). This is not Pierce's domain."
)

h2("World 2: CSA Statistical Testing Layer")
body(
    "Pierce's isolated analysis environment — headline formula testing, format experiments, "
    "platform-specific findings. This is statistical at the article/cluster level, not macro political "
    "trend analysis. The two worlds occasionally reference the same articles but are asking completely "
    "different questions."
)

callout(
    "Don't try to make your testing results match or explain what Justin/Dedra's dashboards show. "
    "They're measuring different things. If someone asks you to reconcile them, the correct answer is: "
    "they're not reconcilable — they're different instruments measuring different phenomena."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 8. THE GOVERNOR SYSTEM — CHRIS'S IMPLICIT APPROVAL
# ═══════════════════════════════════════════════════════════════════════════════
h1("8. The Governor System — Implicit Approval of the 3-Layer Filter")

body(
    "The governor system in the T1 analysis engine has three filter layers:"
)

bullet("Layer 1: Run all analyses — no filter, full statistical output.")
bullet("Layer 2: Filter by relevance to Sarah Price's and Chris's questions — not all findings matter equally; the governor weights headline formula findings highest because that's where Sarah confirmed focus is.")
bullet("Layer 3: Filter out 'obvious to media professional' conclusions — findings that any experienced editor would already know don't belong in the playbook, even if statistically significant.")

body(
    "Chris didn't call this out explicitly, but his description of what he wants from the analysis "
    "environment maps perfectly to this 3-layer architecture. He wants actionable intelligence, not "
    "data dumps. The directional vs. significant distinction (directional → experiments page, significant → playbook rules) "
    "is part of Layer 2/3 filtering."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 9. SWARM RECOMMENDATION AS CSA PRODUCT REQUIREMENT
# ═══════════════════════════════════════════════════════════════════════════════
h1("9. Swarm Recommendation — This Should Be in the PRD as a Product Requirement")

body(
    "Chris described the swarm testing vision as an end goal, not just a nice-to-have. "
    "The sequence he described:"
)

bullet("A story performs above cluster batting average threshold")
bullet("System detects performance (Amplitude/Sigma signal)")
bullet("CSA automatically recommends next cluster topic — 'here's what to write next while this is hot'")
bullet("Writers execute the follow-up article(s)")
bullet("Cluster performance feeds back into the governor to refine future recommendations")

body(
    "He said 'down the line' — but the fact that he described this in enough detail to spec it means "
    "it needs to be in the PRD as a formal product requirement, even if it's in a later phase. "
    "The infrastructure for it (Cluster IDs, Amplitude events, CSA feedback loop) is being built right now."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 10. WHAT CHRIS WANTS FROM PIERCE DIRECTLY
# ═══════════════════════════════════════════════════════════════════════════════
h1("10. What Chris Wants from Pierce — Explicit and Implicit Asks")

h2("Explicit Asks")
bullet("Document the 20-step analysis environment build process so Chris can understand and replicate the sandbox. He asked for this directly — write it up as an onboarding/technical overview doc.")
bullet("Investigate whether Pierce can provision a shared Bitbucket repo space for team collaboration on analysis code. Chris implied this is needed.")
bullet("Build the SEMrush layer as a point-and-click interface for Sarah Price — she should never need to touch the API directly.")
bullet("Own the Gary Tools parameter definition process — drive the editorial taxonomy with Sara Vallone, bring a draft to the meeting next week, then bring the finalized version to Chris for review.")

h2("Implicit Asks (Read Between the Lines)")
bullet(
    "Hold the dev team to pipeline-level thinking. Chris is frustrated that the dev team thinks at "
    "button-placement level. Pierce's job is to hold them to the 3–6 month vision. They're building "
    "half-steps toward an undefined goal. The PRD + the 'control room' model are Pierce's tools for this.",
    bold_prefix="Pipeline thinking:"
)
bullet(
    "Know the batting average metric cold. Chris brings this number to division leadership. "
    "If he ever asks Pierce to pull or verify it, that needs to be a fast answer.",
    bold_prefix="Batting average:"
)
bullet(
    "Don't wait for SEO team review on non-SEO modules. Gary's tool is not primarily an SEO play. "
    "Jim Robinson's novelty concern is real but should not gate the claims validation modules.",
    bold_prefix="SEO gatekeeping:"
)
bullet(
    "Anything that can generalize to the news/L&E division gets greenlit automatically org-wide. "
    "When scoping features or experiments, always ask: can this be applied broadly? If yes, frame it that way.",
    bold_prefix="Auto-greenlight rule:"
)
bullet(
    "The 'control room' is the north star. Trend signals in → content generation → drafts greenlit/revised → "
    "distribution → performance data feeds back. Everything Pierce builds should be traceable to this vision.",
    bold_prefix="Control room model:"
)

# ═══════════════════════════════════════════════════════════════════════════════
# 11. CHRIS'S PRIORITIZATION FRAMEWORK
# ═══════════════════════════════════════════════════════════════════════════════
h1("11. Chris's Prioritization Framework")

body(
    "Chris uses a simple mental model for prioritization:"
)

callout(
    "Economic impact × (inverse of) complexity. Easy wins with good economic return go first. "
    "Hard + low impact = back burner. When in doubt, ask: what's the traffic delta and how much work does it take?"
)

body(
    "Applied to Pierce's current project list, this framework would rank:"
)

bullet("T1 headline formula analysis — very high impact, done. Now the question is: does SEMrush extend it cheaply? (Yes.)")
bullet("Cluster batting average tracking — already instrumented. Keep feeding it.")
bullet("Gary Tools editorial parameters — moderate lift on accuracy/trust, moderate complexity. Chris-directed.")
bullet("Partner content audit (P15) — potentially huge impact, low complexity. Chris controlling the timing.")
bullet("United Robots capture (P12) — high impact, moderate complexity. Revenue share recovery.")
bullet("P7 tracker — low complexity once infra lands, but infra gated on dev. Back burner for now per Chris.")

# ═══════════════════════════════════════════════════════════════════════════════
# 12. DEV TEAM — CHRIS'S VIEW
# ═══════════════════════════════════════════════════════════════════════════════
h1("12. Chris's View of the Dev Team")

body(
    "This is worth understanding clearly — it shapes how you should position your own work in meetings "
    "that include the dev team."
)

bullet("They think at button-placement level, not pipeline level. They are good engineers building features; they are not product thinkers.")
bullet("They're missing the entire workstream that happens after the CSA — distribution, performance tracking, feedback loop. In their mental model, CSA = done when draft is generated.")
bullet("The dev standup is splitting into 2 teams as the upstatement work expands. Cross-training is maintained via shared Slack.")
bullet("The PRD is Pierce's primary tool for forcing pipeline-level thinking. If it's not in the PRD, it won't get built.")
bullet("Rajiv set CSA to 'max effort' mode + increased allowed tokens — a quality improvement signal from the dev side.")

# ═══════════════════════════════════════════════════════════════════════════════
# 13. REUTERS / PARTNER CONTENT EVALUATION
# ═══════════════════════════════════════════════════════════════════════════════
h1("13. Reuters Evaluation — Partner Content Replacement")

body(
    "Reuters is being evaluated as a replacement for the current partner content provider. "
    "This is Kathy's workstream. It is relevant to P15 because if Reuters replaces the current partner, "
    "the inventory optimization exercise will need to account for a different content mix. "
    "Know that this evaluation is in flight before you start any P15 scoping work."
)

# ═══════════════════════════════════════════════════════════════════════════════
# 14. YOUR ACTION ITEMS FROM THIS MEETING
# ═══════════════════════════════════════════════════════════════════════════════
h1("14. Your Action Items from This Meeting")

h2("Immediate")
bullet("Write up the 20-step analysis environment build process — Chris asked explicitly. This is a real deliverable.")
bullet("Investigate Bitbucket provisioning — can Pierce create a shared team repo space?")
bullet("Validate that T1 analysis cross-platform comparisons respect ecosystem boundaries (app-based vs. web-based). Flag and document any findings that commingle them.")

h2("Before Next PRD Revision")
bullet("Add LTV=0 syndication framework as an explicit analytical framing section.")
bullet("Add swarm testing vision as a formal Phase 2/3 product requirement.")
bullet("Add syndication ecosystem taxonomy (app-based captured vs. web-based competitive) as the foundational analytical framework — everything gets filtered through this.")

h2("Medium-Term")
bullet("Partner content audit (P15) — be ready to move fast when Chris gives the green-light. The work is mostly analytical, not technical.")
bullet("Keep the cluster batting average metric current and know how to pull it on demand.")

h2("Watch List (No Action Yet)")
bullet("Reuters evaluation — Kathy's work. Know when it resolves before starting P15.")
bullet("Swarm testing automation — the Cluster ID infra is being built now. When PGS-40 lands, this becomes buildable.")
bullet("Standup split — monitor for any coordination gaps between the two emerging dev teams.")

# ═══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════════════════════════════════
spacer(20)
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = footer_p.add_run("Synthesized from Chris Palo huddle transcript · April 3, 2026 · ops-hub session")
rf.font.size = Pt(8.5)
rf.font.name = "Calibri"
rf.font.color.rgb = GRAY
rf.italic = True

doc.save(OUT)
print(f"Saved: {OUT}")
