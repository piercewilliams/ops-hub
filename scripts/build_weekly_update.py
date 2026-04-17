"""
Populate Sarah Price's National Team Weekly Update template.
Opens the template .docx, fills all fields in-place, saves as new file.
All formatting is inherited from the template — no manual styling needed.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

TEMPLATE = os.path.expanduser('~/Downloads/Template National Team Weekly Update.docx')
OUT      = os.path.expanduser('~/Downloads/National_Team_Weekly_Update_4_7_populated.docx')

doc  = Document(TEMPLATE)
body = doc.element.body

XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

# ── Low-level helpers ──────────────────────────────────────────────────────

def make_run(text, size_pt, color_hex, bold=False):
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    fonts = OxmlElement('w:rFonts')
    fonts.set(qn('w:ascii'), 'Arial')
    fonts.set(qn('w:hAnsi'), 'Arial')
    rPr.append(fonts)
    for tag in ('w:sz', 'w:szCs'):
        el = OxmlElement(tag)
        el.set(qn('w:val'), str(int(size_pt * 2)))
        rPr.append(el)
    col = OxmlElement('w:color')
    col.set(qn('w:val'), color_hex)
    rPr.append(col)
    if bold:
        rPr.append(OxmlElement('w:b'))
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(XML_SPACE, 'preserve')
    t.text = text
    r.append(t)
    return r

def make_linebreak():
    r = OxmlElement('w:r')
    r.append(OxmlElement('w:br'))
    return r

def clear_runs(para_elem):
    for r in list(para_elem):
        if r.tag == qn('w:r'):
            para_elem.remove(r)

def set_para(para_elem, text, size_pt, color_hex, bold=False):
    """Replace all runs in a paragraph with a single run."""
    clear_runs(para_elem)
    para_elem.append(make_run(text, size_pt, color_hex, bold))

def set_cell(cell, run_specs):
    """
    Fill a table cell. run_specs = list of (text, size_pt, color_hex, bold).
    Multiple specs are separated by line breaks within the same cell paragraph.
    """
    p = cell.paragraphs[0]._element
    clear_runs(p)
    for i, (text, size_pt, color_hex, bold) in enumerate(run_specs):
        if i > 0:
            p.append(make_linebreak())
        p.append(make_run(text, size_pt, color_hex, bold))

def make_para_elem(text, size_pt, color_hex, bold=False, bullet=False):
    """
    Create a new w:p element styled with Arial.
    bullet=True applies the template's list bullet formatting (numId=1).
    """
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')

    if bullet:
        # Match template bullet paragraph exactly: numId=1, ilvl=0
        numPr = OxmlElement('w:numPr')
        ilvl  = OxmlElement('w:ilvl');  ilvl.set(qn('w:val'),  '0')
        numId = OxmlElement('w:numId'); numId.set(qn('w:val'), '1')
        numPr.append(ilvl); numPr.append(numId)
        pPr.append(numPr)
        # Indentation matching template bullet: left=560, hanging=280
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'),    '560')
        ind.set(qn('w:hanging'), '280')
        pPr.append(ind)

    # Tight spacing matching template bullet paragraphs: 40/40 twips
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'),   '40')
    spacing.set(qn('w:after'),    '40')
    spacing.set(qn('w:line'),     '240')
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)
    p.append(pPr)
    p.append(make_run(text, size_pt, color_hex, bold))
    return p

def replace_with_paras(placeholder_elem, para_specs):
    """
    Insert new paragraphs in sequence after placeholder, then remove placeholder.
    para_specs = [(text, size_pt, color_hex, bold, bullet), ...]
    """
    cursor = placeholder_elem
    for spec in para_specs:
        text, size_pt, color_hex, bold = spec[0], spec[1], spec[2], spec[3]
        bullet = spec[4] if len(spec) > 4 else False
        new_p = make_para_elem(text, size_pt, color_hex, bold, bullet)
        cursor.addnext(new_p)
        cursor = new_p
    placeholder_elem.getparent().remove(placeholder_elem)

# ── Grab reference elements BEFORE any modifications ──────────────────────
# Body element order from template (verified by inspection):
# [0]  w:p  NATIONAL TEAM WEEKLY UPDATE
# [1]  w:p  [Date]
# [2]  w:tbl stats
# [3]  w:p  CLUSTER PERFORMANCE HIGHLIGHTS
# [4]  w:p  '' (performance placeholder)
# [5]  w:p  CHANGES & DECISIONS
# [6]  w:tbl STOP/FLAG/NEW/TRACKING
# [7]  w:p  ACTIVE TEST PERFORMANCE
# [8]  w:p  '' with bullet (active test placeholder)
# [9]  w:p  CSA & TOOLING UPDATES
# [10] w:p  '' with bullet (CSA placeholder)
# [11] w:p  ACTION ITEMS
# [12] w:tbl action items
# [13] w:p  ''
# [14] w:p  footer

date_para          = body[1]
perf_placeholder   = body[4]
active_placeholder = body[8]
csa_placeholder    = body[10]
footer_para        = body[14]

stats_tbl   = doc.tables[0]
changes_tbl = doc.tables[1]
actions_tbl = doc.tables[2]

# ── 1. Date ───────────────────────────────────────────────────────────────
set_para(date_para, 'April 7, 2026', 9.5, '777777')

# ── 2. Stats table ────────────────────────────────────────────────────────
# Each stats cell has two paragraphs: para[0] = number (empty in template),
# para[1] = label (has placeholder text that must be replaced).

def set_stats_cell(cell, number, label):
    """Fill a two-paragraph stats cell: number on top, label below."""
    p0 = cell.paragraphs[0]._element
    clear_runs(p0)
    p0.append(make_run(number, 13, '2E75B6', False))
    p1 = cell.paragraphs[1]._element
    clear_runs(p1)
    p1.append(make_run(label, 8.5, '555555', False))

def set_stats_cell_single(cell, text):
    """Fill a stats cell that has no number — just a single label line."""
    p0 = cell.paragraphs[0]._element
    clear_runs(p0)                                  # leave top line empty
    p1 = cell.paragraphs[1]._element
    clear_runs(p1)
    p1.append(make_run(text, 8.5, '555555', False))

row0 = stats_tbl.rows[0]
set_stats_cell(row0.cells[0], '366,000+',  'Page Views \u00b7 Apr 1\u20136')
set_stats_cell(row0.cells[1], '347',       'Articles Published')
set_stats_cell(row0.cells[2], '9',         'Properties')
set_stats_cell_single(row0.cells[3],       'YTD Best 6-Day Start')

# ── 3. CLUSTER PERFORMANCE HIGHLIGHTS body ────────────────────────────────
replace_with_paras(perf_placeholder, [
    ('US Weekly and Women\u2019s World top the engagement rankings \u2014 increase publishing on these. Country music and creature features with numbers in headlines performed well across Tier 1 sites.',
     9.5, '000000', False, True),
    ('Miami Herald led as the top swarming parent site (66,483 PVs). Sacramento Bee flagged for expanded testing as a child site in new cluster categories.',
     9.5, '000000', False, True),
    ('Swarming generated 110,142 total PVs across 28 articles \u2014 more than nostalgia and YouTube combined; top cluster (7 Dogs Viral) reached 62,938 PVs.',
     9.5, '000000', False, True),
])

# ── 4. CHANGES & DECISIONS table descriptions ─────────────────────────────
change_descs = [
    "Nostalgia test deprioritized. After 70 articles averaging ~400 PVs each, results were marginal. The team will only pursue nostalgia if a clearly high-potential angle emerges.",
    "US Weekly image captions: move content to the description box (not caption box) per publisher guidance from Emma. The field does not appear on the live page \u2014 likely used for SEO.",
    "Swarming operationalized into a playbook. Trigger: 20,000 clicks within 24 hrs. Top angles: emotional aftermath, debunking/explainer, light viral animal stories.",
    "Two additional CSA personas going live today (total: 3), including SmartNews and Apple News options. Designed to generate more variance than previous versions.",
    "Published links must be added to the tracker after publishing \u2014 critical for US Weekly and Women\u2019s World data collection.",
]
for i, desc in enumerate(change_descs):
    set_cell(changes_tbl.rows[i].cells[1], [(desc, 9.5, '000000', False)])

# ── 5. ACTIVE TEST PERFORMANCE body ───────────────────────────────────────
replace_with_paras(active_placeholder, [
    ('Trigger threshold: 20,000 clicks within the first 24 hours of publication.',
     9.5, '000000', False, True),
    ('Top follow-up angles: emotional aftermath, debunking/explainer framing, light viral animal stories.',
     9.5, '000000', False, True),
    ('Author consistency has been standard; cross-author and weekend coverage testing proposed. Most swarms = 1 follow-up \u2014 worth testing deliberately.',
     9.5, '000000', False, True),
])

# ── 6. CSA & TOOLING UPDATES body ─────────────────────────────────────────
replace_with_paras(csa_placeholder, [
    ('Intros are being made longer. Headline monitoring continues.',
     9.5, '000000', False, True),
    ('~90\u201395% of titles currently generated via Claude directly (not CSA interface).',
     9.5, '000000', False, True),
])

# ── 7. ACTION ITEMS table ─────────────────────────────────────────────────
action_rows = [
    ('Sara Vallone',
     'Add non-CSA headline column to tracker; Share Sam\u2019s US Weekly headline analysis (action verbs / patterns) with full team and Pierce',
     'This week'),
    ('Sara V / Sarah P',
     'Track persona performance; log persona selections and corresponding PV data; swarming playbook',
     'Ongoing'),
    ('The group',
     'Add published links to tracker after every publish \u2014 USW and Women\u2019s World especially critical',
     'Ongoing'),
]
for i, (owner, action, timing) in enumerate(action_rows):
    row = actions_tbl.rows[i + 1]  # row 0 is the header
    set_cell(row.cells[0], [(owner,  9.0, '000000', False)])
    set_cell(row.cells[1], [(action, 9.0, '000000', False)])
    set_cell(row.cells[2], [(timing, 9.0, '000000', False)])

# ── 8. Footer date ────────────────────────────────────────────────────────
for t in footer_para.findall('.//' + qn('w:t')):
    if t.text and 'National Team Weekly' in t.text:
        t.text = 'National Team Weekly  \u00b7  April 7, 2026  \u00b7  Content & Programming'

# ── Save ──────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f'Saved: {OUT}')
